import streamlit as st
import os
import tempfile
from groq import Groq
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io

# ─── PAGINA INSTELLINGEN ───────────────────────────────────────────────
st.set_page_config(
    page_title="MeetingMate",
    page_icon="🎤",
    layout="centered"
)

# ─── WACHTWOORD BEVEILIGING ────────────────────────────────────────────
def check_password():
    if "authenticated" not in st.session_state:
        st.session_state.authenticated = False

    if not st.session_state.authenticated:
        st.markdown("""
            <div style='text-align: center; padding: 50px 0 20px 0;'>
                <h1>🎤 MeetingMate</h1>
                <p style='color: gray;'>Voer het wachtwoord in om door te gaan</p>
            </div>
        """, unsafe_allow_html=True)

        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            password = st.text_input("Wachtwoord", type="password", key="password_input")
            if st.button("Inloggen", use_container_width=True):
                if password == os.environ.get("APP_PASSWORD", "meetingmate123"):
                    st.session_state.authenticated = True
                    st.rerun()
                else:
                    st.error("❌ Verkeerd wachtwoord!")
        return False
    return True

# ─── GROQ TRANSCRIPTIE ─────────────────────────────────────────────────
def transcribe_audio(audio_file):
    client = Groq(api_key=os.environ.get("GROQ_API_KEY"))
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=".m4a") as tmp:
        tmp.write(audio_file.read())
        tmp_path = tmp.name
    
    try:
        with open(tmp_path, "rb") as f:
            transcription = client.audio.transcriptions.create(
                file=(os.path.basename(tmp_path), f.read(), "audio/m4a"),
                model="whisper-large-v3",
                language="nl",
                response_format="text"
            )
        return transcription
    finally:
        os.unlink(tmp_path)

# ─── NOTULEN GENEREREN ─────────────────────────────────────────────────
def generate_notulen(transcriptie, vergadering_info):
    client = Groq(api_key=os.environ.get("GROQ_API_KEY"))
    
    prompt = f"""Je bent een professionele notulist. Maak uitgebreide en gestructureerde notulen op basis van de volgende transcriptie.

Vergadering informatie:
- Naam: {vergadering_info['naam']}
- Datum: {vergadering_info['datum']}
- Aanwezigen: {vergadering_info['aanwezigen']}
- Locatie: {vergadering_info['locatie']}

Transcriptie:
{transcriptie}

Maak de notulen in het Nederlands met de volgende structuur:
1. **Vergaderinformatie** (naam, datum, aanwezigen, locatie)
2. **Samenvatting** (korte samenvatting van de vergadering)
3. **Besproken punten** (alle besproken onderwerpen met details)
4. **Besluiten** (alle genomen besluiten)
5. **Actiepunten** (wie doet wat en wanneer)
6. **Volgende vergadering** (indien besproken)

Schrijf professioneel en volledig."""

    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
        max_tokens=4000
    )
    
    return response.choices[0].message.content

# ─── WORD DOCUMENT MAKEN ───────────────────────────────────────────────
def create_word_document(notulen, vergadering_info):
    doc = Document()
    
    # Titel
    title = doc.add_heading('MeetingMate - Notulen', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Vergadering info
    doc.add_heading(vergadering_info['naam'], level=1)
    
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('Datum:', vergadering_info['datum']),
        ('Locatie:', vergadering_info['locatie']),
        ('Aanwezigen:', vergadering_info['aanwezigen']),
        ('Gegenereerd op:', datetime.now().strftime('%d-%m-%Y %H:%M'))
    ]
    
    for i, (label, value) in enumerate(info_data):
        row = info_table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
    
    doc.add_paragraph()
    
    # Notulen inhoud
    for line in notulen.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('**') and line.endswith('**'):
            p = doc.add_paragraph()
            run = p.add_run(line.strip('**'))
            run.bold = True
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line.strip():
            doc.add_paragraph(line)
    
    # Opslaan in geheugen
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# ─── HOOFDAPP ──────────────────────────────────────────────────────────
def main():
    # Header
    st.markdown("""
        <div style='text-align: center; padding: 20px 0;'>
            <h1>🎤 MeetingMate</h1>
            <p style='color: gray;'>Automatische notulen van jouw vergadering</p>
        </div>
    """, unsafe_allow_html=True)
    
    # Uitloggen knop
    col1, col2, col3 = st.columns([3, 1, 1])
    with col3:
        if st.button("🚪 Uitloggen"):
            st.session_state.authenticated = False
            st.rerun()
    
    st.divider()
    
    # ── STAP 1: Vergadering info ──
    st.subheader("📋 Stap 1: Vergadering informatie")
    
    col1, col2 = st.columns(2)
    with col1:
        vergadering_naam = st.text_input("📌 Naam vergadering", placeholder="Bijv. Teamoverleg Marketing")
        vergadering_datum = st.date_input("📅 Datum", datetime.now())
    with col2:
        vergadering_locatie = st.text_input("📍 Locatie", placeholder="Bijv. Vergaderzaal A / Teams")
        vergadering_aanwezigen = st.text_input("👥 Aanwezigen", placeholder="Bijv. Jan, Lisa, Mohammed")
    
    vergadering_info = {
        'naam': vergadering_naam or "Vergadering",
        'datum': vergadering_datum.strftime('%d-%m-%Y'),
        'locatie': vergadering_locatie or "Niet opgegeven",
        'aanwezigen': vergadering_aanwezigen or "Niet opgegeven"
    }
    
    st.divider()
    
    # ── STAP 2: Audio ──
    st.subheader("🎵 Stap 2: Audio uploaden")
    
    audio_tab1, audio_tab2 = st.tabs(["📁 Bestand uploaden", "🎤 Opnemen"])
    
    audio_data = None
    
    with audio_tab1:
        uploaded_file = st.file_uploader(
            "Upload audio bestand",
            type=["mp3", "mp4", "m4a", "wav", "ogg", "webm"],
            help="Maximaal 25MB"
        )
        if uploaded_file:
            st.audio(uploaded_file)
            audio_data = uploaded_file
    
    with audio_tab2:
        st.info("📱 **Op iPhone:** Gebruik de ingebouwde microfoon hieronder")
        recorded_audio = st.audio_input("🎤 Klik om op te nemen")
        if recorded_audio:
            audio_data = recorded_audio
    
    st.divider()
    
    # ── STAP 3: Verwerken ──
    st.subheader("⚡ Stap 3: Notulen genereren")
    
    if audio_data is None:
        st.warning("⬆️ Upload eerst een audio bestand of neem op")
    else:
        st.success("✅ Audio gereed!")
        
        if st.button("🚀 Genereer Notulen", use_container_width=True, type="primary"):
            
            # Transcriberen
            with st.spinner("🎯 Audio wordt getranscribeerd..."):
                try:
                    transcriptie = transcribe_audio(audio_data)
                    st.session_state.transcriptie = transcriptie
                    st.success("✅ Transcriptie klaar!")
                except Exception as e:
                    st.error(f"❌ Fout bij transcriptie: {str(e)}")
                    return
            
            # Notulen genereren
            with st.spinner("📝 Notulen worden gegenereerd..."):
                try:
                    notulen = generate_notulen(transcriptie, vergadering_info)
                    st.session_state.notulen = notulen
                    st.session_state.vergadering_info = vergadering_info
                    st.success("✅ Notulen gegenereerd!")
                except Exception as e:
                    st.error(f"❌ Fout bij genereren: {str(e)}")
                    return
    
    # ── RESULTATEN ──
    if "notulen" in st.session_state:
        st.divider()
        st.subheader("📄 Notulen")
        
        tab1, tab2 = st.tabs(["📝 Notulen", "🎤 Transcriptie"])
        
        with tab1:
            st.markdown(st.session_state.notulen)
        
        with tab2:
            st.text_area("Volledige transcriptie", st.session_state.transcriptie, height=300)
        
        st.divider()
        
        # ── STAP 4: Downloaden ──
        st.subheader("📥 Stap 4: Notulen downloaden")
        
        word_buffer = create_word_document(
            st.session_state.notulen,
            st.session_state.vergadering_info
        )
        
        st.download_button(
            label="📥 Download als Word",
            data=word_buffer,
            file_name=f"Notulen_{vergadering_naam}_{datetime.now().strftime('%d%m%Y')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

# ─── APP STARTEN ───────────────────────────────────────────────────────
if __name__ == "__main__":
    if check_password():
        main()